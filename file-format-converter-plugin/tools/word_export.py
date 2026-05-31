"""
Word文档导出工具
"""
import os
import tempfile
from collections.abc import Generator
from datetime import datetime
from typing import Any

from dify_plugin import Tool
from dify_plugin.entities.tool import ToolInvokeMessage


class WordExportTool(Tool):
    """Word文档导出工具"""
    
    def _invoke(self, tool_parameters: dict[str, Any]) -> Generator[ToolInvokeMessage]:
        """
        执行工具调用
        
        Args:
            tool_parameters: 工具参数
                - markdown_content: Markdown内容
                - document_name: 文档名称（可选）
                
        Yields:
            ToolInvokeMessage: 工具调用消息
        """
        # 获取参数
        markdown_content = tool_parameters.get('markdown_content', '')
        document_name = tool_parameters.get('document_name', 'document')
        
        if not markdown_content:
            yield self.create_text_message("错误：Markdown内容不能为空")
            return
        
        # 清理文件名
        document_name = self._sanitize_filename(document_name)
        
        # 生成临时文件路径
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"{document_name}_{timestamp}.docx"
        temp_dir = tempfile.gettempdir()
        output_path = os.path.join(temp_dir, filename)
        
        try:
            # 转换为Word文档 - 使用简单的实现
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
            
            doc = Document()
            
            # 简单的markdown解析和转换
            lines = markdown_content.split('\n')
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                
                # 标题
                if line.startswith('# '):
                    p = doc.add_heading(line[2:], level=1)
                elif line.startswith('## '):
                    p = doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    p = doc.add_heading(line[4:], level=3)
                # 列表
                elif line.startswith('- ') or line.startswith('* '):
                    doc.add_paragraph(line[2:], style='List Bullet')
                elif line.startswith('1. ') or line.startswith('2. '):
                    doc.add_paragraph(line[3:], style='List Number')
                # 普通段落
                else:
                    doc.add_paragraph(line)
            
            # 保存文档
            doc.save(output_path)
            
            # 读取文件内容
            with open(output_path, 'rb') as f:
                file_content = f.read()
            
            # 返回文件
            yield self.create_blob_message(
                blob=file_content,
                meta={
                    'mime_type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                    'filename': filename
                }
            )
            yield self.create_text_message(f"✅ Word文档已生成：{filename}")
            
        except Exception as e:
            yield self.create_text_message(f"❌ 生成Word文档失败：{str(e)}")
        
        finally:
            # 清理临时文件
            if os.path.exists(output_path):
                try:
                    os.remove(output_path)
                except:
                    pass
    
    def _sanitize_filename(self, filename: str) -> str:
        """清理文件名，移除非法字符"""
        # 移除文件扩展名
        if filename.endswith('.docx'):
            filename = filename[:-5]
        
        # 移除非法字符
        illegal_chars = '<>:"/\\|?*'
        for char in illegal_chars:
            filename = filename.replace(char, '_')
        
        # 限制长度
        if len(filename) > 100:
            filename = filename[:100]
        
        return filename or 'document'
