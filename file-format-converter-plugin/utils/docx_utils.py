"""
Word文档转换工具
"""
import markdown
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from bs4 import BeautifulSoup
import re


def markdown_to_docx(markdown_content: str, output_path: str) -> str:
    """
    将Markdown内容转换为Word文档
    
    Args:
        markdown_content: Markdown格式的文本内容
        output_path: 输出文件路径
        
    Returns:
        输出文件的路径
    """
    # 将Markdown转换为HTML
    html_content = markdown.markdown(
        markdown_content,
        extensions=['extra', 'codehilite', 'tables', 'fenced_code']
    )
    
    # 解析HTML
    soup = BeautifulSoup(html_content, 'html.parser')
    
    # 创建Word文档
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # 处理HTML元素
    _process_html_elements(soup, doc)
    
    # 保存文档
    doc.save(output_path)
    return output_path


def _process_html_elements(soup, doc):
    """处理HTML元素并添加到Word文档"""
    for element in soup.children:
        if element.name == 'h1':
            _add_heading(doc, element.get_text(), level=1)
        elif element.name == 'h2':
            _add_heading(doc, element.get_text(), level=2)
        elif element.name == 'h3':
            _add_heading(doc, element.get_text(), level=3)
        elif element.name == 'h4':
            _add_heading(doc, element.get_text(), level=4)
        elif element.name == 'p':
            _add_paragraph(doc, element)
        elif element.name == 'ul':
            _add_list(doc, element, ordered=False)
        elif element.name == 'ol':
            _add_list(doc, element, ordered=True)
        elif element.name == 'pre':
            _add_code_block(doc, element.get_text())
        elif element.name == 'blockquote':
            _add_blockquote(doc, element.get_text())
        elif element.name == 'table':
            _add_table(doc, element)
        elif element.name == 'hr':
            doc.add_paragraph('─' * 50)


def _add_heading(doc, text, level=1):
    """添加标题"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT


def _add_paragraph(doc, element):
    """添加段落，支持内联格式"""
    paragraph = doc.add_paragraph()
    
    for child in element.children:
        if isinstance(child, str):
            paragraph.add_run(child)
        elif child.name == 'strong' or child.name == 'b':
            run = paragraph.add_run(child.get_text())
            run.bold = True
        elif child.name == 'em' or child.name == 'i':
            run = paragraph.add_run(child.get_text())
            run.italic = True
        elif child.name == 'code':
            run = paragraph.add_run(child.get_text())
            run.font.name = 'Courier New'
            run.font.color.rgb = RGBColor(220, 50, 47)
        elif child.name == 'a':
            run = paragraph.add_run(child.get_text())
            run.font.color.rgb = RGBColor(0, 0, 255)
            run.underline = True
        else:
            paragraph.add_run(child.get_text())


def _add_list(doc, element, ordered=False):
    """添加列表"""
    for idx, li in enumerate(element.find_all('li', recursive=False)):
        text = li.get_text()
        if ordered:
            text = f"{idx + 1}. {text}"
        else:
            text = f"• {text}"
        doc.add_paragraph(text, style='List Bullet' if not ordered else 'List Number')


def _add_code_block(doc, text):
    """添加代码块"""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    paragraph.paragraph_format.left_indent = Inches(0.5)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(6)


def _add_blockquote(doc, text):
    """添加引用块"""
    paragraph = doc.add_paragraph(text)
    paragraph.paragraph_format.left_indent = Inches(0.5)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.runs[0]
    run.font.italic = True
    run.font.color.rgb = RGBColor(100, 100, 100)


def _add_table(doc, element):
    """添加表格"""
    rows = element.find_all('tr')
    if not rows:
        return
    
    # 获取列数
    cols = len(rows[0].find_all(['th', 'td']))
    
    # 创建表格
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = 'Light Grid Accent 1'
    
    # 填充表格
    for i, row in enumerate(rows):
        cells = row.find_all(['th', 'td'])
        for j, cell in enumerate(cells):
            table.rows[i].cells[j].text = cell.get_text().strip()
            # 表头加粗
            if cell.name == 'th':
                for paragraph in table.rows[i].cells[j].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
